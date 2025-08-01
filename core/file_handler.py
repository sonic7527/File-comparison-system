import os
import pandas as pd
from datetime import datetime
import streamlit as st
from docx import Document
from openpyxl import load_workbook
import re

def get_file_type(filename):
    """根據副檔名獲取文件類型"""
    ext = os.path.splitext(filename)[1].lower()
    if ext in ['.docx', '.doc']:
        return 'docx'
    if ext in ['.xlsx', '.xls']:
        return 'xlsx'
    return 'unknown'

def save_uploaded_file(uploaded_file, directory):
    """儲存上傳的檔案"""
    os.makedirs(directory, exist_ok=True)
    filepath = os.path.join(directory, uploaded_file.name)
    with open(filepath, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return filepath

def parse_excel_fields(excel_path):
    """
    解析Excel欄位，並支援從第三欄的說明中提取下拉選單選項。
    """
    try:
        df = pd.read_excel(excel_path, header=None)
        if df.empty:
            return []

        field_definitions = []
        for index, row in df.iterrows():
            field_name = str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else None
            if not field_name:
                continue

            field_value = str(row.iloc[1]).strip() if len(row) > 1 and pd.notna(row.iloc[1]) else ""
            description = str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else ""
            
            # --- 下拉選單邏輯 ---
            dropdown_options = []
            if "這遠可以做成下拉式選單" in description or "這邊可以做成下拉式選單" in description:
                # 使用正則表達式或字串分割來提取選項
                # 假設選項格式為 "1.選項一 2.選項二" 或 "選項一\n選項二"
                # 清理觸發詞
                clean_desc = description.replace("這遠可以做成下拉式選單", "").replace("這邊可以做成下拉式選單", "").strip()
                
                # 按換行符分割
                options = [opt.strip() for opt in clean_desc.split('\n') if opt.strip()]
                
                # 如果換行符無效，嘗試用數字編號分割
                if len(options) <= 1:
                    # 使用正則表達式尋找 "1. xxx" 或 "1 xxx" 的模式
                    options = re.split(r'\d+\.|\d+\s', clean_desc)
                    options = [opt.strip() for opt in options if opt.strip()]

                if options:
                    dropdown_options = options
            
            field_definitions.append({
                'name': field_name,
                'value': field_value,
                'description': description,
                'dropdown_options': dropdown_options # 新增欄位
            })
        
        return field_definitions
    except Exception as e:
        st.error(f"解析Excel欄位時發生錯誤: {e}")
        return []


def generate_document(template_path, field_values):
    """根據文件類型生成文件"""
    file_type = get_file_type(os.path.basename(template_path))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = "generated_files"
    os.makedirs(output_dir, exist_ok=True)
    
    # 創建獨一無二的檔名
    base_name = os.path.splitext(os.path.basename(template_path))[0]
    output_filename = f"{base_name}_{timestamp}.{file_type}"
    output_path = os.path.join(output_dir, output_filename)

    try:
        if file_type == 'docx':
            doc = Document(template_path)
            # 替換段落和表格中的佔位符
            for para in doc.paragraphs:
                for key, value in field_values.items():
                    para.text = para.text.replace(f"{{{{{key}}}}}", str(value))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for key, value in field_values.items():
                                para.text = para.text.replace(f"{{{{{key}}}}}", str(value))
            doc.save(output_path)
        
        elif file_type == 'xlsx':
            workbook = load_workbook(template_path)
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            for key, value in field_values.items():
                                cell.value = cell.value.replace(f"{{{{{key}}}}}", str(value))
            workbook.save(output_path)
        
        else:
            st.error(f"不支援的檔案類型: {file_type}")
            return None
            
        return output_path
    except Exception as e:
        st.error(f"生成文件時發生錯誤: {e}")
        return None