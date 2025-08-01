import streamlit as st
import pandas as pd
import os
import sqlite3
from typing import Dict, List, Any
from datetime import datetime
import json
from PIL import Image
import io
import base64

class DocumentGenerator:
    def __init__(self):
        self.templates_dir = "data/document_templates"
        self.db_path = "data/document_templates.db"
        self.generated_dir = "generated_files"
        self._ensure_directories()
        self._init_database()
    
    def _ensure_directories(self):
        """確保必要的目錄存在"""
        os.makedirs(self.templates_dir, exist_ok=True)
        os.makedirs(self.generated_dir, exist_ok=True)
    
    def _init_database(self):
        """初始化資料庫 - 支援群組化管理"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # 基本資料群組表
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS field_groups (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        name TEXT NOT NULL,
                        description TEXT,
                        category TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                """)
                
                # 基本資料欄位表
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS field_definitions (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        group_id INTEGER,
                        field_name TEXT NOT NULL,
                        field_description TEXT,
                        field_type TEXT DEFAULT 'text',
                        sample_value TEXT,
                        is_required BOOLEAN DEFAULT 1,
                        display_order INTEGER DEFAULT 0,
                        has_dropdown BOOLEAN DEFAULT 0,
                        dropdown_options TEXT,
                        FOREIGN KEY (group_id) REFERENCES field_groups (id)
                    )
                """)
                
                # 檢查並添加新欄位（用於資料庫遷移）
                cursor.execute("PRAGMA table_info(field_definitions)")
                columns = [column[1] for column in cursor.fetchall()]
                
                if 'has_dropdown' not in columns:
                    cursor.execute("ALTER TABLE field_definitions ADD COLUMN has_dropdown BOOLEAN DEFAULT 0")
                    st.info("✅ 資料庫已更新：添加 has_dropdown 欄位")
                
                if 'dropdown_options' not in columns:
                    cursor.execute("ALTER TABLE field_definitions ADD COLUMN dropdown_options TEXT")
                    st.info("✅ 資料庫已更新：添加 dropdown_options 欄位")
                
                # 範本群組表
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS template_groups (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        field_group_id INTEGER,
                        name TEXT NOT NULL,
                        description TEXT,
                        template_type TEXT,
                        page_size TEXT DEFAULT 'A4',
                        page_orientation TEXT DEFAULT 'PORTRAIT',
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        FOREIGN KEY (field_group_id) REFERENCES field_groups (id)
                    )
                """)
                
                # 範本文件表
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS template_files (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        template_group_id INTEGER,
                        file_name TEXT NOT NULL,
                        file_path TEXT NOT NULL,
                        file_type TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        FOREIGN KEY (template_group_id) REFERENCES template_groups (id)
                    )
                """)
                
                conn.commit()
        except Exception as e:
            st.error(f"❌ 資料庫初始化失敗：{str(e)}")
    
    def create_field_group(self, name: str, description: str, category: str = "一般") -> int:
        """創建基本資料群組"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO field_groups (name, description, category)
                    VALUES (?, ?, ?)
                """, (name, description, category))
                conn.commit()
                return cursor.lastrowid
        except Exception as e:
            st.error(f"❌ 創建基本資料群組失敗：{str(e)}")
            return None
    
    def add_field_definition(self, group_id: int, field_name: str, field_description: str, 
                           field_type: str = "text", sample_value: str = "", 
                           is_required: bool = True, display_order: int = 0,
                           has_dropdown: bool = False, dropdown_options: List[str] = None):
        """添加欄位定義到群組"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # 將下拉選單選項轉換為 JSON 字串
                dropdown_options_json = json.dumps(dropdown_options) if dropdown_options else None
                
                cursor.execute("""
                    INSERT INTO field_definitions 
                    (group_id, field_name, field_description, field_type, sample_value, is_required, display_order,
                     has_dropdown, dropdown_options)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (group_id, field_name, field_description, field_type, sample_value, 
                     is_required, display_order, has_dropdown, dropdown_options_json))
                conn.commit()
                return True
        except Exception as e:
            st.error(f"❌ 添加欄位定義失敗：{str(e)}")
            return False
    
    def get_all_field_groups(self) -> List[Dict]:
        """獲取所有基本資料群組"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT fg.*, COUNT(fd.id) as field_count
                    FROM field_groups fg
                    LEFT JOIN field_definitions fd ON fg.id = fd.group_id
                    GROUP BY fg.id
                    ORDER BY fg.created_at DESC
                """)
                return [dict(row) for row in cursor.fetchall()]
        except Exception as e:
            st.error(f"❌ 獲取基本資料群組失敗：{str(e)}")
            return []
    
    def get_field_group_details(self, group_id: int) -> Dict:
        """獲取基本資料群組詳細資訊"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # 獲取群組資訊
                cursor.execute("SELECT * FROM field_groups WHERE id = ?", (group_id,))
                group = cursor.fetchone()
                if not group:
                    return None
                
                # 獲取欄位定義
                cursor.execute("""
                    SELECT * FROM field_definitions 
                    WHERE group_id = ? 
                    ORDER BY display_order, field_name
                """, (group_id,))
                fields = []
                for row in cursor.fetchall():
                    field_dict = dict(row)
                    
                    # 解析下拉選單選項
                    if field_dict.get('dropdown_options'):
                        try:
                            field_dict['dropdown_options'] = json.loads(field_dict['dropdown_options'])
                        except:
                            field_dict['dropdown_options'] = []
                    else:
                        field_dict['dropdown_options'] = []
                    
                    fields.append(field_dict)
                
                return {
                    'group': dict(group),
                    'fields': fields
                }
        except Exception as e:
            st.error(f"❌ 獲取群組詳細資訊失敗：{str(e)}")
            return None
    
    def create_template_group(self, field_group_id: int, name: str, description: str, 
                            template_type: str, page_size: str = "A4", 
                            orientation: str = "PORTRAIT") -> int:
        """創建範本群組"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO template_groups 
                    (field_group_id, name, description, template_type, page_size, page_orientation)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (field_group_id, name, description, template_type, page_size, orientation))
                conn.commit()
                return cursor.lastrowid
        except Exception as e:
            st.error(f"❌ 創建範本群組失敗：{str(e)}")
            return None
    
    def add_template_file(self, template_group_id: int, file_name: str, file_path: str, file_type: str):
        """添加範本文件"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO template_files 
                    (template_group_id, file_name, file_path, file_type)
                    VALUES (?, ?, ?, ?)
                """, (template_group_id, file_name, file_path, file_type))
                conn.commit()
                return True
        except Exception as e:
            st.error(f"❌ 添加範本文件失敗：{str(e)}")
            return False

    def delete_template_group(self, template_group_id: int) -> bool:
        """刪除範本群組及其相關文件"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # 先獲取範本群組信息
                cursor.execute("SELECT field_group_id FROM template_groups WHERE id = ?", (template_group_id,))
                result = cursor.fetchone()
                if not result:
                    return False
                    
                field_group_id = result[0]
                
                # 獲取相關文件路徑並刪除物理文件
                cursor.execute("SELECT file_path FROM template_files WHERE template_group_id = ?", (template_group_id,))
                files = cursor.fetchall()
                
                for (file_path,) in files:
                    full_path = os.path.join(self.templates_dir, file_path)
                    if os.path.exists(full_path):
                        os.remove(full_path)
                
                # 刪除資料庫記錄
                cursor.execute("DELETE FROM template_files WHERE template_group_id = ?", (template_group_id,))
                cursor.execute("DELETE FROM template_groups WHERE id = ?", (template_group_id,))
                
                # 檢查是否還有其他範本群組使用相同的 field_group
                cursor.execute("SELECT COUNT(*) FROM template_groups WHERE field_group_id = ?", (field_group_id,))
                count = cursor.fetchone()[0]
                
                # 如果沒有其他範本群組使用，也刪除 field_group 和相關 field_definitions
                if count == 0:
                    cursor.execute("DELETE FROM field_definitions WHERE group_id = ?", (field_group_id,))
                    cursor.execute("DELETE FROM field_groups WHERE id = ?", (field_group_id,))
                
                conn.commit()
                return True
        except Exception as e:
            st.error(f"❌ 刪除範本群組失敗：{str(e)}")
            return False
    
    def get_all_template_groups(self) -> List[Dict]:
        """獲取所有範本群組"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT tg.*, fg.name as field_group_name, fg.description as field_group_description,
                           COUNT(tf.id) as file_count
                    FROM template_groups tg
                    LEFT JOIN field_groups fg ON tg.field_group_id = fg.id
                    LEFT JOIN template_files tf ON tg.id = tf.template_group_id
                    GROUP BY tg.id
                    ORDER BY tg.created_at DESC
                """)
                return [dict(row) for row in cursor.fetchall()]
        except Exception as e:
            st.error(f"❌ 獲取範本群組失敗：{str(e)}")
            return []
    
    def get_template_group_details(self, template_group_id: int) -> Dict:
        """獲取範本群組詳細資訊"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # 獲取範本群組資訊
                cursor.execute("""
                    SELECT tg.*, fg.name as field_group_name, fg.description as field_group_description
                    FROM template_groups tg
                    LEFT JOIN field_groups fg ON tg.field_group_id = fg.id
                    WHERE tg.id = ?
                """, (template_group_id,))
                template_group = cursor.fetchone()
                if not template_group:
                    return None
                
                # 獲取範本文件
                cursor.execute("""
                    SELECT * FROM template_files 
                    WHERE template_group_id = ?
                    ORDER BY created_at DESC
                """, (template_group_id,))
                files = [dict(row) for row in cursor.fetchall()]
                
                # 獲取對應的基本資料群組
                field_group_details = self.get_field_group_details(template_group['field_group_id'])
                
                return {
                    'template_group': dict(template_group),
                    'files': files,
                    'field_group': field_group_details
                }
        except Exception as e:
            st.error(f"❌ 獲取範本群組詳細資訊失敗：{str(e)}")
            return None
    
    def convert_image_to_excel(self, image_path: str, output_path: str, page_size: str = "A4", orientation: str = "PORTRAIT"):
        """將圖片轉換為A4格式的Excel文件"""
        try:
            st.info("✅ 開始轉換圖片為Excel格式")
            
            # 讀取圖片
            with Image.open(image_path) as img:
                # 獲取圖片尺寸
                img_width, img_height = img.size
                
                # A4尺寸設定（以點為單位）
                if orientation == "LANDSCAPE":
                    page_width = 29.7 * 28.35  # A4橫式寬度
                    page_height = 21.0 * 28.35  # A4橫式高度
                else:
                    page_width = 21.0 * 28.35   # A4直式寬度
                    page_height = 29.7 * 28.35  # A4直式高度
                
                # 計算縮放比例
                scale_x = page_width / img_width
                scale_y = page_height / img_height
                scale = min(scale_x, scale_y)  # 保持比例
                
                # 調整圖片尺寸
                new_width = int(img_width * scale)
                new_height = int(img_height * scale)
                img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # 創建Excel文件
                from openpyxl import Workbook
                from openpyxl.drawing.image import Image as XLImage
                
                wb = Workbook()
                ws = wb.active
                
                # 設定頁面大小
                ws.page_setup.paperSize = ws.page_setup.PAPERSIZE_A4
                if orientation == "LANDSCAPE":
                    ws.page_setup.orientation = ws.page_setup.ORIENTATION_LANDSCAPE
                
                # 將圖片保存為臨時文件
                temp_img_path = f"temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                img_resized.save(temp_img_path)
                
                # 插入圖片到Excel
                xl_img = XLImage(temp_img_path)
                ws.add_image(xl_img, 'A1')
                
                # 調整欄寬和行高以適應圖片
                ws.column_dimensions['A'].width = new_width / 7
                ws.row_dimensions[1].height = new_height * 0.75
                
                # 保存Excel文件
                wb.save(output_path)
                
                # 清理臨時文件
                os.remove(temp_img_path)
                
                st.success(f"✅ 圖片已轉換為A4 {orientation}格式的Excel文件")
                return output_path
                
        except Exception as e:
            st.error(f"❌ 圖片轉Excel失敗：{str(e)}")
            return None
    
    def convert_image_to_word(self, image_path: str, output_path: str, page_size: str = "A4", orientation: str = "PORTRAIT"):
        """將圖片轉換為A4格式的Word文件"""
        try:
            st.info("✅ 開始轉換圖片為Word格式")
            
            from docx import Document
            from docx.shared import Inches, Cm
            from docx.enum.section import WD_ORIENT
            
            # 創建Word文件
            doc = Document()
            
            # 設定A4頁面大小
            section = doc.sections[0]
            if orientation == "LANDSCAPE":
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Cm(29.7)
                section.page_height = Cm(21.0)
            else:
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
            
            # 讀取圖片
            with Image.open(image_path) as img:
                # 獲取圖片尺寸
                img_width, img_height = img.size
                
                # 計算A4尺寸（以英寸為單位）
                if orientation == "LANDSCAPE":
                    page_width_inch = 29.7 / 2.54
                    page_height_inch = 21.0 / 2.54
                else:
                    page_width_inch = 21.0 / 2.54
                    page_height_inch = 29.7 / 2.54
                
                # 計算縮放比例
                scale_x = page_width_inch / (img_width / 96)
                scale_y = page_height_inch / (img_height / 96)
                scale = min(scale_x, scale_y)
                
                # 調整圖片尺寸
                new_width = img_width * scale
                new_height = img_height * scale
                
                # 將圖片保存為臨時文件
                temp_img_path = f"temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                img_resized = img.resize((int(new_width), int(new_height)), Image.Resampling.LANCZOS)
                img_resized.save(temp_img_path)
                
                # 插入圖片到Word
                doc.add_picture(temp_img_path, width=Inches(new_width / 96), height=Inches(new_height / 96))
                
                # 清理臨時文件
                os.remove(temp_img_path)
            
            # 保存Word文件
            doc.save(output_path)
            
            st.success(f"✅ 圖片已轉換為A4 {orientation}格式的Word文件")
            return output_path
            
        except Exception as e:
            st.error(f"❌ 圖片轉Word失敗：{str(e)}")
            return None
    
    def convert_image_to_pdf(self, image_path: str, output_path: str, page_size: str = "A4", orientation: str = "PORTRAIT"):
        """將圖片轉換為A4格式的PDF文件"""
        try:
            st.info("✅ 開始轉換圖片為PDF格式")
            
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4, A4_LANDSCAPE
            from reportlab.lib.units import cm
            
            # 選擇頁面大小
            if orientation == "LANDSCAPE":
                pagesize = A4_LANDSCAPE
            else:
                pagesize = A4
            
            # 創建PDF文件
            c = canvas.Canvas(output_path, pagesize=pagesize)
            
            # 讀取圖片
            with Image.open(image_path) as img:
                # 獲取圖片尺寸
                img_width, img_height = img.size
                
                # 計算A4尺寸（以點為單位）
                if orientation == "LANDSCAPE":
                    page_width = 29.7 * cm
                    page_height = 21.0 * cm
                else:
                    page_width = 21.0 * cm
                    page_height = 29.7 * cm
                
                # 計算縮放比例
                scale_x = page_width / img_width
                scale_y = page_height / img_height
                scale = min(scale_x, scale_y)
                
                # 調整圖片尺寸
                new_width = img_width * scale
                new_height = img_height * scale
                
                # 計算居中位置
                x_offset = (page_width - new_width) / 2
                y_offset = (page_height - new_height) / 2
                
                # 將圖片保存為臨時文件
                temp_img_path = f"temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                img_resized = img.resize((int(new_width), int(new_height)), Image.Resampling.LANCZOS)
                img_resized.save(temp_img_path)
                
                # 插入圖片到PDF
                c.drawImage(temp_img_path, x_offset, y_offset, width=new_width, height=new_height)
                
                # 清理臨時文件
                os.remove(temp_img_path)
            
            # 保存PDF文件
            c.save()
            
            st.success(f"✅ 圖片已轉換為A4 {orientation}格式的PDF文件")
            return output_path
            
        except Exception as e:
            st.error(f"❌ 圖片轉PDF失敗：{str(e)}")
            return None
    
    def generate_document_from_template(self, template_group_id: int, input_data: Dict, output_name: str, output_format: str = None, selected_file_index: int = 0) -> str:
        """從範本群組生成文件"""
        try:
            template_details = self.get_template_group_details(template_group_id)
            if not template_details:
                st.error("❌ 找不到範本群組資訊")
                return None
            
            template_group = template_details['template_group']
            files = template_details['files']
            
            if not files:
                st.error("❌ 範本群組中沒有文件")
                return None
            
            # 檢查選擇的文件索引是否有效
            if selected_file_index >= len(files):
                st.error(f"❌ 選擇的文件索引超出範圍")
                return None
            
            # 使用選擇的文件作為範本
            template_file = files[selected_file_index]
            template_path = os.path.join(self.templates_dir, template_file['file_path'])
            
            if not os.path.exists(template_path):
                st.error(f"❌ 範本文件不存在：{template_path}")
                return None
            
            # 如果指定了輸出格式，使用指定的格式；否則使用範本原始格式
            if output_format:
                if output_format.lower() == 'excel':
                    return self._generate_excel_document(template_path, input_data, output_name)
                elif output_format.lower() == 'word':
                    return self._generate_word_document(template_path, input_data, output_name)
                elif output_format.lower() == 'pdf':
                    return self._generate_pdf_document(template_path, input_data, output_name)
                else:
                    st.error(f"❌ 不支援的輸出格式：{output_format}")
                    return None
            else:
            # 根據文件類型處理
                file_type = template_file.get('file_type', 'Unknown').lower()
                if file_type == 'excel':
                return self._generate_excel_document(template_path, input_data, output_name)
                elif file_type == 'word':
                return self._generate_word_document(template_path, input_data, output_name)
                elif file_type == 'pdf':
                    return self._generate_pdf_document(template_path, input_data, output_name)
            else:
                    st.error(f"❌ 不支援的文件類型：{file_type}")
                return None
                
        except Exception as e:
            st.error(f"❌ 生成文件失敗：{str(e)}")
            return None
    
    def _generate_excel_document(self, template_path: str, input_data: Dict, output_name: str) -> str:
        """生成Excel文件"""
        try:
            st.info("✅ 複製Excel模板")
            
            import shutil
            output_path = os.path.join(self.generated_dir, f"{output_name}.xlsx")
            shutil.copy2(template_path, output_path)
            
            # 使用openpyxl填入資料
            try:
                from openpyxl import load_workbook
                workbook = load_workbook(output_path, data_only=False, keep_vba=False)
                
                filled_count = 0
                for sheet in workbook.sheetnames:
                    sheet_obj = workbook[sheet]
                    st.info(f"✅ 處理工作表：{sheet}")
                    
                    for row in sheet_obj.iter_rows():
                        for cell in row:
                            if cell.value:
                                cell_value_str = str(cell.value).strip()
                                # 檢查是否在輸入資料中有對應的欄位
                                for field_name, value in input_data.items():
                                    if cell_value_str == field_name.strip():
                                        cell.value = value
                                        filled_count += 1
                                        st.info(f"✅ 填入 {field_name}: {value}")
                                        break
                
                workbook.save(output_path)
                st.success(f"✅ 成功填入 {filled_count} 個欄位")
                return output_path
                
            except Exception as e:
                st.warning(f"⚠️ openpyxl處理失敗：{str(e)}")
                return None
                
        except Exception as e:
            st.error(f"❌ Excel生成失敗：{str(e)}")
            return None
    
    def _generate_word_document(self, template_path: str, input_data: Dict, output_name: str) -> str:
        """生成Word文件"""
        try:
            st.info("✅ 複製Word模板")
            
            import shutil
            output_path = os.path.join(self.generated_dir, f"{output_name}.docx")
            shutil.copy2(template_path, output_path)
            
            # 使用python-docx填入資料
            try:
                from docx import Document
                doc = Document(output_path)
                
                filled_count = 0
                for paragraph in doc.paragraphs:
                    for field_name, value in input_data.items():
                        if field_name in paragraph.text:
                            paragraph.text = paragraph.text.replace(field_name, str(value))
                            filled_count += 1
                            st.info(f"✅ 填入 {field_name}: {value}")
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for field_name, value in input_data.items():
                                if field_name in cell.text:
                                    cell.text = cell.text.replace(field_name, str(value))
                                    filled_count += 1
                                    st.info(f"✅ 填入 {field_name}: {value}")
                
                doc.save(output_path)
                st.success(f"✅ 成功填入 {filled_count} 個欄位")
                return output_path
                
            except ImportError:
                st.error("❌ 需要安裝 python-docx：pip install python-docx")
                return None
            except Exception as e:
                st.warning(f"⚠️ Word處理失敗：{str(e)}")
                return None
                
        except Exception as e:
            st.error(f"❌ Word生成失敗：{str(e)}")
            return None

    def get_file_properties(self, file):
        """自動判斷檔案類型和方向"""
        file_extension = os.path.splitext(file.name)[1].lower()
        
        # 判斷文件類型
        if file_extension in ['.xlsx', '.xls']:
            template_type = 'Excel'
        elif file_extension in ['.docx', '.doc']:
            template_type = 'Word'
        elif file_extension == '.pdf':
            template_type = 'PDF'
        elif file.type.startswith('image/'):
            template_type = 'Image' # 標記為圖片，後續會轉換
        else:
            template_type = 'Unknown'
            
        # 判斷方向 (目前主要針對圖片和PDF)
        orientation = 'PORTRAIT' # 預設為直式
        try:
            if template_type in ['Image', 'PDF']:
                file_bytes = io.BytesIO(file.getvalue())
                if template_type == 'Image':
                    with Image.open(file_bytes) as img:
                        width, height = img.size
                        if width > height:
                            orientation = 'LANDSCAPE'
                elif template_type == 'PDF':
                    from PyPDF2 import PdfReader
                    reader = PdfReader(file_bytes)
                    page = reader.pages[0]
                    box = page.mediabox
                    if box.width > box.height:
                        orientation = 'LANDSCAPE'
        except Exception as e:
            st.warning(f"⚠️ 無法自動判斷檔案 {file.name} 的方向，將使用預設值。錯誤：{e}")

        return template_type, orientation

    def _generate_pdf_document(self, template_path: str, input_data: Dict, output_name: str) -> str:
        """生成PDF文件"""
        try:
            st.info("✅ 複製PDF模板")
            
            import shutil
            output_path = os.path.join(self.generated_dir, f"{output_name}.pdf")
            shutil.copy2(template_path, output_path)
            
            # 使用PyPDF2填入資料
            try:
                from PyPDF2 import PdfReader, PdfWriter
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                from io import BytesIO
                
                # 讀取原始PDF
                reader = PdfReader(template_path)
                writer = PdfWriter()
                
                filled_count = 0
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    
                    # 創建一個新的PDF頁面來覆蓋原始頁面
                    packet = BytesIO()
                    can = canvas.Canvas(packet, pagesize=A4)
                    
                    # 在頁面上填入資料
                    y_position = 800  # 起始Y位置
                    for field_name, value in input_data.items():
                        # 簡單的文字填入（實際應用中可能需要更複雜的定位邏輯）
                        can.drawString(100, y_position, f"{field_name}: {value}")
                        y_position -= 20
                        filled_count += 1
                        st.info(f"✅ 填入 {field_name}: {value}")
                    
                    can.save()
                    packet.seek(0)
                    
                    # 將新頁面與原始頁面合併
                    overlay = PdfReader(packet)
                    page.merge_page(overlay.pages[0])
                    writer.add_page(page)
                
                # 保存新的PDF
                with open(output_path, "wb") as output_file:
                    writer.write(output_file)
                
                st.success(f"✅ 成功填入 {filled_count} 個欄位")
                return output_path
                
            except ImportError:
                st.error("❌ 需要安裝 PyPDF2：pip install PyPDF2")
                return None
            except Exception as e:
                st.warning(f"⚠️ PDF處理失敗：{str(e)}")
                return None
                
        except Exception as e:
            st.error(f"❌ PDF生成失敗：{str(e)}")
            return None

    def parse_excel_field_definitions(self, excel_file) -> List[Dict]:
        """解析上傳的Excel檔案，自動創建欄位定義"""
        try:
            # 讀取Excel檔案
            df = pd.read_excel(excel_file)
            
            # 檢查是否為三欄格式
            if len(df.columns) < 3:
                st.error("❌ Excel檔案格式錯誤：需要至少3欄（欄位名稱、輸入內容、說明）")
                return []
            
            # 假設前三欄分別為：欄位名稱、輸入內容、說明
            field_name_col = df.columns[0]
            input_content_col = df.columns[1]
            description_col = df.columns[2]
            
            field_definitions = []
            
            for index, row in df.iterrows():
                field_name = str(row[field_name_col]).strip()
                input_content = str(row[input_content_col]).strip()
                description = str(row[description_col]).strip()
                
                # 跳過空行
                if not field_name or field_name == 'nan':
                    continue
                
                # 判斷欄位類型
                field_type = self._determine_field_type(input_content, description)
                
                # 判斷是否需要下拉選單
                has_dropdown = self._has_dropdown_options(description)
                dropdown_options = self._extract_dropdown_options(description) if has_dropdown else []
                
                # 判斷是否為必填欄位
                is_required = bool(input_content.strip()) and input_content != 'nan'
                
                field_definitions.append({
                    'field_name': field_name,
                    'field_description': description,
                    'field_type': field_type,
                    'sample_value': input_content if input_content != 'nan' else '',
                    'has_dropdown': has_dropdown,
                    'dropdown_options': dropdown_options,
                    'is_required': is_required,  # 如果輸入內容為空或'nan'，則為選填
                    'display_order': index
                })
            
            return field_definitions
            
        except Exception as e:
            st.error(f"❌ 解析Excel檔案失敗：{str(e)}")
            return []
    
    def _determine_field_type(self, input_content: str, description: str) -> str:
        """判斷欄位類型"""
        # 清理輸入內容，移除 'nan' 值
        if input_content == 'nan' or not input_content:
            input_content = ''
        
        # 根據欄位名稱和說明優先判斷
        field_lower = description.lower()
        content_lower = input_content.lower()
        
        # 優先根據欄位名稱判斷
        if '案名' in description or '名稱' in description or '地點' in description or '容量' in description:
            return 'text'
        
        # 檢查是否為電話號碼（優先檢查欄位名稱）
        if '電話' in description:
            return 'phone'
        
        # 檢查是否為電子郵件
        if '@' in input_content:
            return 'email'
        
        # 檢查是否為日期格式
        date_patterns = ['/', '-', '年', '月', '日']
        if any(pattern in input_content for pattern in date_patterns):
            return 'date'
        
        # 數字判斷 - 只有當欄位名稱明確表示是數字時才判斷為數字
        if input_content and input_content.replace('.', '').replace('-', '').replace('+', '').replace(',', '').isdigit():
            # 如果欄位說明中包含明確的數字相關詞彙
            number_keywords = ['數量', '金額', '價格', '費用', '率', '比例', '百分比']
            if any(keyword in description for keyword in number_keywords):
                if '.' in input_content:
                    return 'number'  # 小數
        else:
                    return 'number'  # 整數
            # 電話號碼特別處理 - 長度在8-15位的純數字
            elif '電話' in description and len(input_content.replace('-', '').replace(' ', '')) >= 8:
                return 'phone'
            else:
                # 其他情況預設為文字，避免誤判
                return 'text'
        
        # 預設為文字
        return 'text'
    
    def _has_dropdown_options(self, description: str) -> bool:
        """檢查是否有下拉選單選項"""
        dropdown_indicators = ['下拉式選單', '下拉選單', '選單', '選項', '是否', '躉售']
        return any(indicator in description for indicator in dropdown_indicators)
    
    def _extract_dropdown_options(self, description: str) -> List[str]:
        """從說明中提取下拉選單選項"""
        options = []
        
        # 尋找數字開頭的選項（如：1.選項1, 2.選項2）
        import re
        pattern = r'(\d+\.\s*[^,\n]+)'
        matches = re.findall(pattern, description)
        
        for match in matches:
            # 移除數字和點號
            option = re.sub(r'^\d+\.\s*', '', match).strip()
            if option:
                options.append(option)
        
        # 如果沒有找到數字開頭的選項，檢查是否有「是否」或「躉售」類型的選項
        if not options and ('是否' in description or '躉售' in description):
            # 尋找常見的是否選項
            yes_no_patterns = ['是/否', '是|否', '是、否', '是 否']
            for pattern in yes_no_patterns:
                if pattern in description:
                    parts = pattern.replace('|', '/').replace('、', '/').split('/')
                    options.extend([part.strip() for part in parts if part.strip()])
                    break
            
            # 如果還是沒有找到，使用預設的是否選項
            if not options:
                options = ['是', '否']
        
        return options

def document_generator_tab():
    """文件生成頁面 - 群組化管理版本"""
    st.header("📄 智能文件生成系統")
    
    generator = DocumentGenerator()
    
    # 創建標籤頁
    tab1, tab2, tab3 = st.tabs(["📤 範本群組", "📝 生成文件", "📊 管理面板"])
    
    with tab1:
        st.subheader("📤 範本群組管理")
        st.markdown("""
        **概念說明：**
        - **範本群組**：基於基本資料群組創建的不同格式範本
        - **文件格式**：可以是Excel、Word或圖片轉換的文件
        - **A4格式**：自動轉換為A4直式或橫式格式
        """)
        
        # 創建新範本群組
        with st.expander("➕ 創建新的範本群組"):
            st.markdown("""
            **新的工作流程：**
            1. 上傳基本資料 Excel 檔案（三欄格式）
            2. 上傳完成後的表格範本
            3. 系統自動分析並創建範本群組
            """)
            
            with st.form("create_template_group_from_excel"):
                template_name = st.text_input("範本群組名稱", help="例如：台電送件審查範本")
                    template_description = st.text_area("範本說明", help="描述這個範本的用途")
                    
                st.markdown("---")
                
                st.markdown("#### **步驟 1：上傳基本資料**")
                excel_file = st.file_uploader(
                    "上傳基本資料 Excel 檔案",
                    type=['xlsx', 'xls'],
                    help="上傳包含三欄格式的 Excel 檔案（欄位名稱、輸入內容、說明）"
                )
                
                st.markdown("---")
                
                st.markdown("#### **步驟 2：上傳範本檔案**")
                template_files = st.file_uploader(
                    "上傳範本檔案 (可多選)",
                    type=['xlsx', 'xls', 'docx', 'doc', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'bmp'],
                    help="支援所有範本格式，系統會自動判斷文件類型與方向。",
                    accept_multiple_files=True
                            )
                    
                    if st.form_submit_button("📤 創建範本群組"):
                    if template_name and excel_file and template_files:
                        # 自動判斷第一個範本的屬性
                        first_file = template_files[0]
                        template_type, page_orientation = generator.get_file_properties(first_file)

                        if template_type == 'Unknown':
                            st.error(f"❌ 無法識別的檔案類型：{first_file.name}")
                            return

                        # 解析 Excel 檔案
                        field_definitions = generator.parse_excel_field_definitions(excel_file)
                        
                        if field_definitions:
                            # 創建基本資料群組
                            group_name = f"{template_name}_基本資料"
                            group_description = f"基於 {template_name} 的基本資料欄位"
                            field_group_id = generator.create_field_group(group_name, group_description, "範本群組")
                            
                            # 添加欄位定義
                            for field_def in field_definitions:
                                generator.add_field_definition(
                                    field_group_id, 
                                    field_def['field_name'], 
                                    field_def['field_description'],
                                    field_def['field_type'], 
                                    field_def['sample_value'], 
                                    field_def['is_required'], 
                                    field_def['display_order'],
                                    field_def.get('has_dropdown', False),
                                    field_def.get('dropdown_options', [])
                                )
                            
                            # 根據第一個檔案的類型創建範本群組
                            # 注意：如果上傳的是圖片，類型會被記錄為最終轉換成的PDF格式
                            final_template_type = 'PDF' if template_type == 'Image' else template_type
                            template_group_id = generator.create_template_group(
                                field_group_id, template_name, template_description,
                                final_template_type, "A4", page_orientation
                            )
                            
                            if template_group_id:
                                # 處理多個文件上傳
                                success_count = 0
                                for template_file in template_files:
                                file_timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                                    current_file_type, current_orientation = generator.get_file_properties(template_file)
                                
                                    if current_file_type == 'Image':
                                        st.info(f"✅ 檢測到圖片文件 {template_file.name}，正在轉換為PDF...")
                                    
                                        temp_img_path = os.path.join(generator.templates_dir, f"temp_{file_timestamp}_{template_file.name}")
                                    with open(temp_img_path, "wb") as f:
                                        f.write(template_file.getvalue())
                                    
                                        final_path = os.path.join(generator.templates_dir, f"template_{file_timestamp}_{template_file.name}.pdf")
                                        success = generator.convert_image_to_pdf(temp_img_path, final_path, "A4", current_orientation)
                                        
                                    os.remove(temp_img_path)
                                    
                                    if success:
                                        file_name = os.path.basename(final_path)
                                            generator.add_template_file(template_group_id, file_name, file_name, 'PDF')
                                            success_count += 1
                                    elif current_file_type != 'Unknown':
                                    # 直接保存文件
                                        original_filename, file_extension = os.path.splitext(template_file.name)
                                        file_name = f"template_{file_timestamp}_{original_filename}{file_extension}"
                                    final_path = os.path.join(generator.templates_dir, file_name)
                                    
                                    with open(final_path, "wb") as f:
                                        f.write(template_file.getvalue())
                                    
                                        generator.add_template_file(template_group_id, file_name, file_name, current_file_type)
                                        success_count += 1
                                
                                if success_count == len(template_files):
                                    st.success(f"✅ 範本群組 '{template_name}' 創建成功，並已成功上傳 {success_count} 個範本檔案！")
                        else:
                                    st.warning(f"⚠️ 範本群組 '{template_name}' 已創建，但只有 {success_count} / {len(template_files)} 個檔案成功上傳。")
                        else:
                            st.warning("⚠️ 請填寫所有必要欄位並上傳檔案")
        
        # 顯示現有範本群組
        st.subheader("📤 現有範本群組")
        template_groups = generator.get_all_template_groups()
        
        if template_groups:
            for group in template_groups:
                with st.expander(f"📤 {group['name']} ({group['file_count']} 個文件)"):
                    col1, col2 = st.columns([4, 1])
                    
                    with col1:
                    st.write(f"**說明：** {group['description']}")
                    st.write(f"**基本資料群組：** {group['field_group_name']}")
                    st.write(f"**文件類型：** {group['template_type']}")
                    st.write(f"**頁面方向：** {'直式' if group['page_orientation'] == 'PORTRAIT' else '橫式'}")
                    st.write(f"**創建時間：** {group['created_at']}")
                    
                    with col2:
                        if st.button(f"🗑️ 刪除", key=f"delete_{group['id']}", type="secondary"):
                            if st.session_state.get(f"confirm_delete_{group['id']}", False):
                                # 執行刪除
                                if generator.delete_template_group(group['id']):
                                    st.success(f"✅ 範本群組 '{group['name']}' 已刪除！")
                                    st.session_state[f"confirm_delete_{group['id']}"] = False
                                    st.experimental_rerun()
                                else:
                                    st.error("❌ 刪除失敗")
                            else:
                                # 顯示確認
                                st.session_state[f"confirm_delete_{group['id']}"] = True
                                st.experimental_rerun()
                        
                        # 顯示確認訊息
                        if st.session_state.get(f"confirm_delete_{group['id']}", False):
                            st.warning("⚠️ 確定要刪除嗎？")
                            col_yes, col_no = st.columns(2)
                            with col_yes:
                                if st.button("確定", key=f"yes_{group['id']}", type="primary"):
                                    if generator.delete_template_group(group['id']):
                                        st.success(f"✅ 範本群組 '{group['name']}' 已刪除！")
                                        st.session_state[f"confirm_delete_{group['id']}"] = False
                                        st.experimental_rerun()
                            with col_no:
                                if st.button("取消", key=f"no_{group['id']}"):
                                    st.session_state[f"confirm_delete_{group['id']}"] = False
                                    st.experimental_rerun()
        else:
            st.info("📝 還沒有創建任何範本群組")
    
    with tab2:
        st.subheader("📝 生成文件")
        st.markdown("""
        **使用流程：**
        1. 選擇範本群組
        2. 系統自動載入對應的基本資料輸入頁面
        3. 填入新資料
        4. 生成新文件
        """)
        
        template_groups = generator.get_all_template_groups()
        if not template_groups:
            st.warning("⚠️ 請先創建範本群組")
            return
        
        # 選擇範本群組
        template_options = {f"{g['name']} - {g['field_group_name']}": g['id'] for g in template_groups}
        selected_template_name = st.selectbox(
            "選擇範本群組",
            list(template_options.keys()),
            help="選擇要使用的範本群組"
        )
        
        if selected_template_name:
            template_group_id = template_options[selected_template_name]
            
            # 獲取範本詳細資訊
            template_details = generator.get_template_group_details(template_group_id)
            
            if template_details and template_details['field_group']:
                field_group = template_details['field_group']
                fields = field_group['fields']
                
                st.success(f"✅ 載入基本資料群組：{field_group['group']['name']}")
                st.info(f"📋 發現 {len(fields)} 個輸入欄位")
                
                # 顯示輸入欄位
                st.subheader("📝 輸入資料")
                
                input_data = {}
                for i, field in enumerate(fields):
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        # 檢查是否有下拉選單選項
                        has_dropdown = field.get('has_dropdown', False)
                        dropdown_options = field.get('dropdown_options', [])
                        
                        if has_dropdown and dropdown_options:
                            # 使用下拉選單
                            value = st.selectbox(
                                f"{field['field_name']}",
                                options=dropdown_options,
                                index=0 if dropdown_options else 0,
                                key=f"dropdown_{i}",
                                help=field['field_description']
                            )
                        elif field['field_type'] == 'number':
                            value = st.number_input(
                                f"{field['field_name']}",
                                value=float(field['sample_value']) if field['sample_value'] and str(field['sample_value']).replace('.', '').isdigit() else 0.0,
                                key=f"input_{i}",
                                help=field['field_description']
                            )
                        elif field['field_type'] == 'date':
                            value = st.date_input(
                                f"{field['field_name']}",
                                key=f"input_{i}",
                                help=field['field_description']
                            )
                        elif field['field_type'] == 'phone':
                            value = st.text_input(
                                f"{field['field_name']}",
                                value=str(field['sample_value']) if field['sample_value'] else "",
                                key=f"input_{i}",
                                help=field['field_description']
                            )
                        else:
                            value = st.text_input(
                                f"{field['field_name']}",
                                value=str(field['sample_value']) if field['sample_value'] else "",
                                key=f"input_{i}",
                                help=field['field_description']
                            )
                    
                    with col2:
                        st.write(f"**說明：** {field['field_description']}")
                        if has_dropdown and dropdown_options:
                            st.write("**下拉選單**")
                        if field['is_required']:
                            st.write("**必填**")
                    
                    input_data[field['field_name']] = value
                
                # 生成文件
                st.subheader("🚀 生成文件")
                
                # 選擇特定的模板文件
                template_files = template_details.get('files', [])
                if template_files:
                    st.info(f"📁 此範本群組包含 {len(template_files)} 個模板文件")
                    
                    # 創建模板文件選項
                    template_file_options = []
                    for i, file_info in enumerate(template_files):
                        file_name = file_info.get('file_name', f'模板文件 {i+1}')
                        template_file_options.append({
                            'index': i,
                            'name': file_name,
                            'file_path': file_info.get('file_path', ''),
                            'file_type': file_info.get('file_type', 'Unknown')
                        })
                    
                    # 顯示模板文件選擇
                    selected_file_index = st.selectbox(
                        "選擇要使用的模板文件",
                        options=range(len(template_file_options)),
                        format_func=lambda x: template_file_options[x]['name'],
                        help="選擇要生成文件的模板文件"
                    )
                    
                    selected_template_file = template_file_options[selected_file_index]
                    st.success(f"✅ 已選擇模板文件：{selected_template_file['name']}")
                else:
                    st.warning("⚠️ 此範本群組沒有模板文件")
                    return
                
                col1, col2 = st.columns(2)
                with col1:
                    output_name = st.text_input(
                        "輸出文件名稱",
                        value=f"生成文件_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                        help="不需要輸入副檔名"
                    )
                
                with col2:
                    # 獲取範本類型
                    template_type = selected_template_file['file_type']
                    
                    # 格式選擇
                    output_format = st.selectbox(
                        "輸出格式",
                        options=['Excel', 'Word', 'PDF'],
                        index=0 if template_type.lower() == 'excel' else (1 if template_type.lower() == 'word' else 2),
                        help="選擇生成文件的格式"
                    )
                
                    if st.button("🚀 生成文件", type="primary"):
                        if output_name:
                            with st.spinner("正在生成文件..."):
                            output_path = generator.generate_document_from_template(template_group_id, input_data, output_name, output_format, selected_file_index)
                                
                                if output_path and os.path.exists(output_path):
                                    st.success("✅ 文件生成成功！")
                                    
                                    # 提供下載連結
                                    with open(output_path, "rb") as f:
                                        file_data = f.read()
                                    
                                    file_extension = os.path.splitext(output_path)[1]
                                if file_extension == ".xlsx":
                                    mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                elif file_extension == ".docx":
                                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                elif file_extension == ".pdf":
                                    mime_type = "application/pdf"
                                else:
                                    mime_type = "application/octet-stream"
                                    
                                    st.download_button(
                                        label="📥 下載文件",
                                        data=file_data,
                                        file_name=f"{output_name}{file_extension}",
                                        mime=mime_type
                                    )
                                    
                                    st.info(f"📁 文件已儲存至：{output_path}")
                                else:
                                    st.error("❌ 文件生成失敗")
                        else:
                            st.warning("⚠️ 請輸入輸出文件名稱")
            else:
                st.warning("⚠️ 無法載入範本詳細資訊")
    
    with tab3:
        st.subheader("📊 管理面板")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric("📋 基本資料群組", len(generator.get_all_field_groups()))
            st.metric("📤 範本群組", len(generator.get_all_template_groups()))
        
        with col2:
            st.metric("📄 生成文件", len(os.listdir(generator.generated_dir)) if os.path.exists(generator.generated_dir) else 0)
            st.metric("✅ 系統狀態", "正常運行")

if __name__ == "__main__":
    document_generator_tab() 